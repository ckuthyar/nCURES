import pandas as pd
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def reader(file):
    df1 = pd.read_excel(file, usecols='A', nrows=23)
    c1= df1.iloc[5:23, 0].tolist()
    
    df2 = pd.read_excel(file, usecols='B', nrows=23)
    c2= df2.iloc[4:23, 0].tolist()
    
    df3 = pd.read_excel(file, usecols='D', nrows=23)
    c3= df3.iloc[5:23, 0].tolist()
    
    df4 = pd.read_excel(file, usecols="B", nrows=2)
    c4 = df4.iloc[0:2,0].tolist()
    
    
    doc = Document()
    para=doc.add_heading("COLOR DOPPLER & 2D ECHOCARDIOGRAPHY (TRANS-THORACIC)",level = 1)
    para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"BSA:{c4[1]:.2f} M². BP: mmHg. Quality of acoustic window: satisfactory.")

    table = doc.add_table(rows=1, cols=2)

    table.style = 'Table Grid'

    col_width = Inches(3.0)
    table.columns[0].width = col_width
    table.columns[1].width = col_width
    left_cell =table.cell(0,0)

    paragraph = left_cell.add_paragraph()
    run = paragraph.add_run("Left atrium:")
    run.bold = True
    run.underline = True

    left_cell.add_paragraph(c2[0]+str(c3[0]))
    left_cell.add_paragraph(c2[1]+str(c3[1]))
    paragraph = left_cell.add_paragraph()
    run = paragraph.add_run("Left ventrials:")
    run.bold = True
    run.underline= True

    left_cell.add_paragraph(c2[2]+str(c3[2]))
    left_cell.add_paragraph(c2[3]+str(c3[3]))
    left_cell.add_paragraph(c2[4]+str(c3[4]))
    left_cell.add_paragraph(c2[5]+str(c3[5]))
    left_cell.add_paragraph("Ejection fraction:[1] %")
    left_cell.add_paragraph("Wall motion abnormality: [2]")
    left_cell.add_paragraph("Diastolic function (assessed by online calculator using 2D, flow Doppler and tissue Doppler measurements and observations):")
    paragraph = left_cell.add_paragraph()
    run = paragraph.add_run("Right ventricles:")
    run.bold = True
    run.underline = True
    left_cell.add_paragraph(c2[6]+str(c3[6]))
    left_cell.add_paragraph(c2[7]+str(c3[7]))
    left_cell.add_paragraph(c2[8]+str(c3[8]))
    left_cell.add_paragraph("RV funtion: [3],TASPE:[4]mm")
    paragraph = left_cell.add_paragraph()
    run = paragraph.add_run("Right strium:")
    run.bold =True
    run.underline = True
    left_cell.add_paragraph(c2[9]+str(c3[9]))
    paragraph = left_cell.add_paragraph()
    run = paragraph.add_run("Mitral valve:")
    run.bold =True
    run.underline = True
    left_cell.add_paragraph("Morphology: [5] ")
    left_cell.add_paragraph("MVA: [6]")
    left_cell.add_paragraph("Flow velocity:")
    left_cell.add_paragraph("   E:[7] cm/sec")
    left_cell.add_paragraph("   A:[8] cm/sec")
    left_cell.add_paragraph("Function: [9]")

    right_cell = table.cell(0,1)

    para = right_cell.add_paragraph()
    run = para.add_run("Aortic Valve:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("Morphology: [10] ")
    right_cell.add_paragraph("MACS: [11]")
    right_cell.add_paragraph("Flow velocity:[12] cm/sec")
    right_cell.add_paragraph("Function: [13] ")
    para = right_cell.add_paragraph()
    run = para.add_run("Tricuspid Valve:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("Morphology: [14]")
    right_cell.add_paragraph("Flow velocity: [15] cm/sec")
    right_cell.add_paragraph("Function: [16] ")
    para = right_cell.add_paragraph()
    run = para.add_run("Pulmonary VAlve:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("Morphology: [17]")
    right_cell.add_paragraph("Flow velocity: [18]cm/sec")
    right_cell.add_paragraph("Function: [19] ")
    para = right_cell.add_paragraph()
    run = para.add_run("IVS:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("[20]Appears intact.")
    para = right_cell.add_paragraph()
    run = para.add_run("IAS:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("[21]Appears intact ")
    para = right_cell.add_paragraph()
    run = para.add_run("Pericardium:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("[22]")
    para = right_cell.add_paragraph()
    run = para.add_run("Endocardium:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("[23]")
    para = right_cell.add_paragraph()
    run = para.add_run("Aorta:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("Aortic Root:[24] mm")
    right_cell.add_paragraph("Arch and DA: [25]")
    para = right_cell.add_paragraph()
    run = para.add_run("Pulmonary artery:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph(c2[14]+str(c3[14]))
    right_cell.add_paragraph("PASP:[26]")
    para = right_cell.add_paragraph()
    run = para.add_run("SVC and IVC:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("[27]")
    para = right_cell.add_paragraph()
    run = para.add_run("Pulmonary veins:")
    run.bold = True
    run.underline = True
    right_cell.add_paragraph("[28]")

    doc.add_heading('IMPRESSION: ', level=2)
    doc.add_paragraph()
    paragraph=doc.add_paragraph()
    run = paragraph.add_run("                                                          Dr. Naveen M. Ballal, MBBS DMRD, Consultant Radiologist")
    run.bold=True
    run.underline = True
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph=doc.add_paragraph()
    run=paragraph.add_run("Note: This study has certain limitations. Depending upon the clinical requirement, further evaluation or follow up may be required, to confirm above findings and to look for abnormalities if any which would have gone undetected in this study.")
    run.font.size = Pt(8)
    file_path=os.getcwd()
    print(file_path)
    
    doc.save("Doc1")
    file_name="Doc1"
    print("Ready")
    return file_name